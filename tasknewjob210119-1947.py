"""
Description: auxiliary scripts
    document = Document('Задание.docx')
    all_paragraphs = document.paragraphs
    desc_task = document.paragraphs
    for p in desc_task:
        print(p.text.encode('utf-8'))
    with open("Пример ответа.json", "r") as read_file:
        data_example_answer = json.load(read_file)
    print('EXAMPLE ANSWER:', data_example_answer)
    print('MY RESULT END :', my_result_end)

EXAMPLE_CLASS
class AnswerToQuestion(Dict):
	''' AnswerOnQuestion '''
	def __init__(self, data):
		Dict.__init__().super()
		name = 'name_model'
		model = data
		
	def __str__(self):
		return self.name

print(
    f'count years: {count_years}',
    f"of leap years are: {response_leap}",
    sep='\n',
    )
"""
import json
import os
from addict import Dict
from docx import Document
import pandas as pd
import datetime
import calendar

path = os.getcwd()
path_files = os.walk(__file__)

with open("DataSet.json", "r") as read_file:
    data = json.load(read_file)

addi_data = Dict(data)
answer = dict()

answer_keys = [
	'result', 
	'client', 
	'rules', 
	'arrayLiability',
	]
all_clients_id = list(addi_data.client.keys())
my_result = dict()
name_credit_history = ['В', 'G', 'N', 'E']
name_rules = [
	'rule_1', 
	'rule_2', 
	'rule_3', 
	'rule_4', 
	'rule_5',
	]

current_year = datetime.datetime.now().strftime('%Y')
min_age = 18
max_age = 55

years = range((int(current_year) - 18), int(current_year))
count_years = len(years)

def find_leap_year(years):
    answer_by_y = 0

    for year in years:
        if calendar.isleap(int(year)):
            answer_by_y = answer_by_y + 1

    return answer_by_y

response_leap = find_leap_year(years)


for i, iclient in enumerate(all_clients_id, 1):
    countOpenLiability = 0
    countCloseLiability = 0
    rules_by_client = []  # для напосления списка правил, сработавших по клиенту
    array_liability = []  # для напосления списка правил, сработавших по клиенту
    
    
    surname = (addi_data.client[iclient].name).split(' ')
    surname_birthday = '|'.join([surname[0], addi_data.client[iclient].birthday])
    
    app_no = addi_data.dataLiability.appNO[iclient]  # список обязательст
    
    date_start = addi_data.dataLiability.dateStart[iclient]  # список дат начала обязательств
    date_end = addi_data.dataLiability.dateEnd[iclient]  # список дат окончания обязательств
    sum_delay = 0 # sum(addi_data.dataLiability.sumDelay[iclient])  # сумма просрочки
    
    for ia, app in enumerate(app_no):
    	if surname_birthday == addi_data.dataLiability.appOwner[iclient][ia]:
    		array_liability.append(app)
    		
    		if date_end[ia] == '' :
    			countOpenLiability += 1
    			sum_delay += addi_data.dataLiability.sumDelay[iclient][ia]
    		elif date_end[ia] != '':
	    		sy = int(date_start[ia][-4:])
	    		sm = int(date_start[ia][3:5])
	    		sd = int(date_start[ia][:2])
	    		
	    		ey = int(date_start[ia][-4:])
	    		em = int(date_end[ia][3:5])
	    		ed = int(date_end[ia][:2])
    			
    			if datetime.datetime(sy, sm, sd) < datetime.datetime(ey, em, ed) or date_end[ia] != '':
    				countCloseLiability += 1
    
    
    
    statusCHIP = addi_data.client[iclient]['citizenship']  # definition of citizenship
    if statusCHIP != 'RU':
        rules_by_client.append(name_rules[0])
        
    
    # DEFINITION AGE
    birthday_ = addi_data.client[iclient].birthday[-4:]
    bdm = addi_data.client[iclient].birthday[3:5]
    bdd = addi_data.client[iclient].birthday[:2]
    age = int(current_year) - int(birthday_)
    age_full = ((datetime.date.today() - \
        datetime.date(int(birthday_), int(bdm), int(bdd))).days - response_leap*366)/365 + response_leap
    if age_full < min_age:
        rules_by_client.append(name_rules[1])
    if age_full > max_age:
        rules_by_client.append(name_rules[2])
        

    # оценка кредитная история
    if sum_delay > 1000:
        credit_history = name_credit_history[0]  # ПЛОХАЯ
    elif sum_delay == 0 and countOpenLiability != 0 or countCloseLiability != 0:
        credit_history = name_credit_history[1]  # ХОРОШАЯ
    elif 0 < sum_delay <= 1000 and countCloseLiability != 0:
        credit_history = name_credit_history[2]  # СРЕДНЯЯ
    else:  # elif app_no == []:
        credit_history = name_credit_history[3]  # ОТСУТСТВУЕТ
        
                
    statusERGIP = addi_data.source.egripStatus[iclient]  # definition of status ERGIP
    if int(statusERGIP) == 0:
        rules_by_client.append(name_rules[3])
        
        
    # Term of business activity
    dateREGIST = addi_data.source.dateRegistration[iclient]  # definition of activity
    current_date = datetime.datetime.now().strftime('%d.%m.%Y')
    date_reg = datetime.datetime(int(dateREGIST[-4:]), int(dateREGIST[3:5]), int(dateREGIST[:2]))
    dateDELTA = (datetime.datetime.now() - date_reg).days
    if int(dateDELTA) < 180:
        rules_by_client.append(name_rules[4])
        
    
    my_result.update(
        {f'client_{i}': {
            # 'surname_birthday': surname_birthday,
            # 'age_full': age_full,
            'rules': rules_by_client,
            'arrayLiability': array_liability,  # list(set(app_no)),
            'countOpenLiability': countOpenLiability,
            'countCloseLiability': countCloseLiability,
            'sumDelays': sum_delay,
            'statusCH': credit_history,
         }}
        )
        

my_result_end = dict({'result': {'client': my_result}})
print('MY RESULT END:\n', my_result_end)
