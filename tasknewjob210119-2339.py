"""
Description: auxiliary scripts
    document = Document('Задание.docx')
    all_paragraphs = document.paragraphs
    desc_task = document.paragraphs
    for p in desc_task:
        print(p.text.encode('utf-8'))
    with open("Пример ответа.json", "r") as read_file:
        data_example_answer = json.load(read_file)
    print('EXAMPLE ANSWER:', data_example_answer)
    print('MY RESULT END :', my_result_end)

EXAMPLE_CLASS
class AnswerToQuestion(Dict):
	''' AnswerOnQuestion '''
	def __init__(self, data):
		Dict.__init__().super()
		name = 'name_model'
		model = data
		
	def __str__(self):
		return self.name

print(
    f'count years: {count_years}',
    f"of leap years are: {response_leap}",
    sep='\n',
    )
"""
import json
import os
from addict import Dict
from docx import Document
import pandas as pd
import datetime
import calendar
from data_to import data_to_resp

path = os.getcwd()
path_files = os.walk(__file__)


def find_leap_year(age):
    """"""
    answer_by_y = 0
    years = range((int(current_year) - int(age)), int(current_year))
    count_years = len(years)
    for year in years:
        if calendar.isleap(int(year)):
            answer_by_y = answer_by_y + 1
    return answer_by_y, count_years


with open("DataSet.json", "r") as read_file:
    data = json.load(read_file)

addi_data = Dict(data_to_resp)
# answer = dict()

answer_keys = [
	'result', 
	'client', 
	'rules', 
	'arrayLiability',
	]

all_clients_id = list(addi_data.client.keys())

my_result = dict()

name_credit_history = ['В', 'G', 'N', 'E']

name_rules = [
	'rule_1', 
	'rule_2', 
	'rule_3', 
	'rule_4', 
	'rule_5',
	]

min_age = 18
max_age = 55
check_status = 'RU'
check_days_reg = 180

current_year = datetime.datetime.now().strftime('%Y')
today = datetime.date.today()

for i, iclient in enumerate(all_clients_id, 1):
    countOpenLiability = 0
    countCloseLiability = 0
    rules_by_client = []  # to fill in the list of rules that worked for the client
    array_liability = []  # to fill in the list of the client's unique liability
    
    surname = (addi_data.client[iclient].name).split(' ')
    surname_birthday = '|'.join([surname[0], addi_data.client[iclient].birthday])
    
    app_no = addi_data.dataLiability.appNO[iclient]  # list liability
    
    date_start = addi_data.dataLiability.dateStart[iclient]  # список дат начала обязательств
    date_end = addi_data.dataLiability.dateEnd[iclient]  # список дат окончания обязательств
    sum_delay = 0 # sum(addi_data.dataLiability.sumDelay[iclient])  # сумма просрочки
    
    for ia, app in enumerate(app_no):
    	if surname_birthday == addi_data.dataLiability.appOwner[iclient][ia]:
    		array_liability.append(app)
    		
    		if date_end[ia] == '' :
    			countOpenLiability += 1
    			sum_delay += addi_data.dataLiability.sumDelay[iclient][ia]
    		elif date_end[ia] != '':
	    		sy = int(date_start[ia][-4:])
	    		sm = int(date_start[ia][3:5])
	    		sd = int(date_start[ia][:2])
	    		
	    		ey = int(date_start[ia][-4:])
	    		em = int(date_end[ia][3:5])
	    		ed = int(date_end[ia][:2])
    			
    			if datetime.datetime(sy, sm, sd) < datetime.datetime(ey, em, ed) or date_end[ia] != '':
    				countCloseLiability += 1    

    # DEFINITION / DETERMINATION OF THE RESIDENT STATUS
    statusCHIP = addi_data.client[iclient]['citizenship']
    if statusCHIP != check_status:
        rules_by_client.append(name_rules[0])        
    
    # DEFINITION / DETERMINING THE CLIENT'S AGE
    birthday_ = addi_data.client[iclient].birthday[-4:]
    bdm = addi_data.client[iclient].birthday[3:5]
    bdd = addi_data.client[iclient].birthday[:2]
    
    age = int(current_year) - int(birthday_)
    convert_birthday = datetime.date(int(birthday_), int(bdm), int(bdd))
    age_full_days = (today - convert_birthday).days
    
    response_leap = find_leap_year(age)[0]
    age_full_year = (age_full_days - response_leap*366)/365 + response_leap
    
    if age_full_year < min_age:
        rules_by_client.append(name_rules[1])
    if age_full_year > max_age:
        rules_by_client.append(name_rules[2])        

    # DEFINITION / ASSESSMENT OF CREDIT HISTORY
    if sum_delay > 1000:
        credit_history = name_credit_history[0]  # BAD
    elif sum_delay == 0 and countOpenLiability != 0 or countCloseLiability != 0:
        credit_history = name_credit_history[1]  # GOOD
    elif 0 < sum_delay <= 1000 and countCloseLiability != 0:
        credit_history = name_credit_history[2]  # MIDDLE
    else:  # elif app_no == []:
        credit_history = name_credit_history[3]  # EMPTY

    # DEFINITION / DETERMINIG THE REGISTRATION STATUS
    statusERGIP = addi_data.source.egripStatus[iclient]  # definition of status ERGIP
    if int(statusERGIP) == 0:
        rules_by_client.append(name_rules[3])        
        
    # DEFINITION / TERM OF BUSINESS ACTIVITY
    dateREGIST = addi_data.source.dateRegistration[iclient]
    current_date = datetime.datetime.now().strftime('%d.%m.%Y')
    try:
        date_reg = datetime.datetime(int(dateREGIST[-4:]), int(dateREGIST[3:5]), int(dateREGIST[:2]))
        dateDELTA = (datetime.datetime.now() - date_reg).days
        if int(dateDELTA) < check_days_reg:
            rules_by_client.append(name_rules[4])        
    except TypeError:
        pass

    # FILLING IN THE RESULTS
    my_result.update(
        {f'client_{i}': {
            # 'surname_birthday': surname_birthday,
            # 'response_leap': response_leap,
            # 'age_full': age_full_year,
            # 'response_count_leap': response_count_leap,
            'rules': rules_by_client,
            'countOpenLiability': countOpenLiability,
            'countCloseLiability': countCloseLiability,
            'sumDelays': sum_delay,
            'statusCH': credit_history,
            'arrayLiability': array_liability,  # list(set(app_no)),
         }}
        )        

my_result_end = dict({'result': {'client': my_result}})
print('MY RESULT END:\n', my_result_end)
