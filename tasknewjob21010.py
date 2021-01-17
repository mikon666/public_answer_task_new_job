"""
Description: auxiliary scripts
    document = Document('Задание.docx')
    all_paragraphs = document.paragraphs
    desc_task = document.paragraphs
    for p in desc_task:
        print(p.text.encode('utf-8'))
    with open("Пример ответа.json", "r") as read_file:
        data_example_answer = json.load(read_file)
    print('EXAMPLE ANSWER:\n', data_example_answer)
    print('MY RESULT END:\n', my_result_end)
"""
import json
import os
from addict import Dict
from docx import Document
import pandas as pd
import datetime

path = os.getcwd()
path_files = os.walk(__file__)

with open("DataSet.json", "r") as read_file:
    data = json.load(read_file)
print(data)
addi_data = Dict(data)
answer = dict()

all_clients_id = list(addi_data.client.keys())
my_result = dict()
name_credit_history = ['E', 'G', 'средняя', 'N']
name_rules = ['rule_1', 'rule_2', 'rule_3', 'rule_4', 'rule_5']
min_age = 18
max_age = 55


for i, iclient in enumerate(all_clients_id, 1):
    countOpenLiability = 0
    countCloseLiability = 0
    rules_by_client = []  # для напосления списка правил, сработавших по клиенту
    
    app_no = addi_data.dataLiability.appNO[iclient]  # список обязательст
    date_start = addi_data.dataLiability.dateStart[iclient]  # список дат начала обязательств
    date_end = addi_data.dataLiability.dateEnd[iclient]  # список дат окончания обязательств
    sum_delay = sum(addi_data.dataLiability.sumDelay[iclient])  # сумма просрочки
    
    statusCHIP = addi_data.client[iclient]['citizenship']  # definition of citizenship
    if statusCHIP != 'РФ':
        rules_by_client.append(name_rules[0])

    for j, d_end in enumerate(date_end):
        if d_end == '':
            countOpenLiability += 1
        if d_end != '' and int(date_start[j][-4:]) < int(d_end[-4:]) \
           and int(d_end[3:5]) < int(date_start[j][3:5]) \
           and int(d_end[:2]) < int(date_start[j][:2]):
            countCloseLiability += 1
            
    birthday_ = addi_data.client[iclient].birthday[-4:]
    current_year = datetime.datetime.now().strftime('%Y')
    age = int(current_year) - int(birthday_)
    if age < min_age:
        rules_by_client.append(name_rules[1])
    if age > max_age:
        rules_by_client.append(name_rules[2])

    # оценка кредитная история
    if sum_delay > 1000:
        credit_history = name_credit_history[0]
    elif sum_delay == 0 and countOpenLiability != 0 or countCloseLiability != 0:
        credit_history = name_credit_history[1]
    elif 0 < sum_delay <= 1000 and countCloseLiability != 0:
        credit_history = name_credit_history[2]
    elif app_no == []:
        credit_history = name_credit_history[3]
                
    statusERGIP = addi_data.source.stopList[iclient]  # definition of status ERGIP
    if int(statusERGIP) == 0:
        rules_by_client.append(name_rules[3])
        
    # Term of business activity
    dateREGIST = addi_data.source.dateRegistration[iclient]  # definition of activity
    current_date = datetime.datetime.now().strftime('%d.%m.%Y')
    date_reg = datetime.datetime(int(dateREGIST[-4:]), int(dateREGIST[3:5]), int(dateREGIST[:2]))
    dateDELTA = (datetime.datetime.now() - date_reg).days
    if int(dateDELTA) < 180:
        rules_by_client.append(name_rules[4])
    
    my_result.update(
        {f'client_{i}': {
            'rules': rules_by_client,
            'arrayLiability': list(set(app_no)),
            'countOpenLiability': countOpenLiability,
            'countCloseLiability': countCloseLiability,
            'sumDelays': sum_delay,
            'statusCH': credit_history,
         }}
        )    

my_result_end = dict({'result': {'client': my_result}})
print('MY RESULT END:\n', my_result_end)

